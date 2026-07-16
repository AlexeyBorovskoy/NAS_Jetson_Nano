"""Correct unsupported Habr publication claims in the bilingual project DOCX."""

from pathlib import Path

from docx import Document


DOCX = Path("docs/articles/publication/NAS_Jetson_Nano_HOME_CLOUD_PROJECT_PROMPTS_RU_EN.docx")

REPLACEMENTS = {
    "Статья опубликована в Habr Sandbox и на GitHub Pages.": (
        "GitHub Pages доступен; статья для Habr не опубликована и остаётся черновиком."
    ),
    "Статья опубликована в Habr Sandbox и на GitHub Pages. На 16.07.2026 обновлены данные VPS, zram, containers и Immich; исторические screenshots получили отдельные date captions.": (
        "GitHub Pages доступен; статья для Habr не опубликована и остаётся черновиком. "
        "На 16.07.2026 отдельные данные VPS, zram, контейнеров и Immich были записаны как датированный снимок; без повторной проверки они не являются текущими."
    ),
    "Published article materials use only redacted screenshots and placeholders.": (
        "Prepared article materials use only redacted screenshots and placeholders; this does not imply publication."
    ),
    "The article was published and later updated with current infrastructure facts while preserving screenshots as historical evidence.": (
        "The article was not published. Its draft was updated with dated infrastructure observations while preserving screenshots as historical evidence; those observations are not current without a new check."
    ),
}


def replace_in_paragraph(paragraph):
    original = paragraph.text
    updated = original
    for old, new in REPLACEMENTS.items():
        updated = updated.replace(old, new)
    if updated != original:
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = updated
        else:
            paragraph.add_run(updated)
        return 1
    return 0


def iter_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def main():
    document = Document(DOCX)
    changed = sum(replace_in_paragraph(p) for p in iter_paragraphs(document))
    if changed == 0:
        raise SystemExit("No matching claims found; document was not changed")
    document.save(DOCX)
    print(f"Corrected {changed} paragraphs in {DOCX}")


if __name__ == "__main__":
    main()
